from docx import Document
import uuid
import os

def save_to_doc(text, output_dir="output"):
    os.makedirs(output_dir, exist_ok=True)
    file_name = f"{uuid.uuid4()}.docx"
    doc = Document()
    doc.add_heading("Generated Maintenance Instructions", level=1)
    doc.add_paragraph(text)
    path = os.path.join(output_dir, file_name)
    doc.save(path)
    return path
